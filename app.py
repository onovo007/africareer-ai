import os
from dotenv import load_dotenv
load_dotenv()
import streamlit as st
import pinecone
from pinecone import Pinecone, ServerlessSpec
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_core.messages import HumanMessage, SystemMessage
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
from datetime import datetime
import json
import hashlib
from collections import Counter
import pandas as pd
import httpx
from urllib.parse import quote_plus, urlparse
import faulthandler
faulthandler.enable()  # on a native crash (exit 139), print the faulting Python frame to the logs

# ----- BRANDING -----
COMPANY_NAME = "Quantium Insights LLC"
APP_NAME = "AfriCareer AI"
TAGLINE = "Intelligent career and academic guidance for African youth and professionals"

# Supported languages - rendered at the TOP of the main page (mobile-first: no sidebar needed)
LANGUAGES = {
    "🇬🇧 English": "English",
    "🇫🇷 Français": "French",
    "🇰🇪 Kiswahili": "Swahili",
    "🇸🇦 العربية": "Arabic",
    "🇳🇬 Hausa": "Hausa",
    "🇳🇬 Pidgin": "Nigerian Pidgin",
    "🇵🇹 Português": "Portuguese",
    "🇪🇸 Español": "Spanish",
    "🇪🇹 አማርኛ": "Amharic",
}
target_lang = "English"  # module-level default; overwritten by the picker in main()

# Countries for sign-in (self-reported -> clean geographic analytics, privacy-friendly)
COUNTRIES = [
    "Select your country",
    # Africa
    "Algeria", "Angola", "Benin", "Botswana", "Burkina Faso", "Burundi", "Cabo Verde",
    "Cameroon", "Central African Republic", "Chad", "Comoros", "Congo (Brazzaville)",
    "Congo (DRC)", "Cote d'Ivoire", "Djibouti", "Egypt", "Equatorial Guinea", "Eritrea",
    "Eswatini", "Ethiopia", "Gabon", "Gambia", "Ghana", "Guinea", "Guinea-Bissau", "Kenya",
    "Lesotho", "Liberia", "Libya", "Madagascar", "Malawi", "Mali", "Mauritania", "Mauritius",
    "Morocco", "Mozambique", "Namibia", "Niger", "Nigeria", "Rwanda", "Sao Tome and Principe",
    "Senegal", "Seychelles", "Sierra Leone", "Somalia", "South Africa", "South Sudan", "Sudan",
    "Tanzania", "Togo", "Tunisia", "Uganda", "Zambia", "Zimbabwe",
    # Rest of world (common)
    "United States", "United Kingdom", "Canada", "Australia", "New Zealand", "Ireland",
    "Germany", "France", "Spain", "Portugal", "Italy", "Netherlands", "Belgium", "Switzerland",
    "Sweden", "Norway", "Denmark", "Finland", "Austria", "Poland", "Greece", "Turkey",
    "China", "India", "Japan", "South Korea", "Singapore", "Malaysia", "Indonesia",
    "Philippines", "Pakistan", "Bangladesh", "United Arab Emirates", "Saudi Arabia", "Qatar",
    "Brazil", "Mexico", "Argentina", "Russia",
    "Other",
]

# Admin credentials (override via env/secret in production; defaults kept for local dev)
ADMIN_USERNAME = os.getenv("ADMIN_USERNAME", "admin")
ADMIN_PASSWORD_HASH = hashlib.sha256(
    os.getenv("ADMIN_PASSWORD", "africareer2024").encode()
).hexdigest()

# ----- ACCESS CONTROL & COST GUARDRAILS -----
# Optional invite code for pilot gating. If unset, the app is open (local/dev).
APP_ACCESS_CODE = os.getenv("APP_ACCESS_CODE", "").strip()
# Cap AI calls per browser session to protect the OpenAI budget on public deployments.
MAX_LLM_CALLS_PER_SESSION = int(os.getenv("MAX_LLM_CALLS_PER_SESSION", "30"))
QUOTA_MESSAGE = (
    "⚠️ You've reached this session's usage limit for AI guidance. "
    "This limit keeps the service free and available for everyone. "
    "Please refresh the page later to continue, or contact the AfriCareer AI team for extended access."
)

# ===== SAFETY GUARDRAILS SYSTEM MESSAGE - MAXIMALLY HELPFUL =====
SAFETY_SYSTEM_MESSAGE = """You are AfriCareer AI, a comprehensive career guidance assistant for African youth.

YOUR MISSION: Help with ANY career, job, education, or professional development question. Be as helpful as possible!

YOU SHOULD ANSWER ALL QUESTIONS ABOUT:
✅ Career guidance, planning, and transitions (any industry, any level)
✅ Resume/CV creation, improvement, and tailoring (for any job)
✅ Job searching strategies and platforms
✅ Interview preparation (for ANY specific job or company)
✅ Interview tips, common questions, how to answer behavioral questions
✅ Salary negotiation and job offer evaluation
✅ Educational pathways, courses, certifications, degrees
✅ Professional development and skill building
✅ Entrepreneurship, business planning, startups
✅ Workplace issues, conflicts, communication
✅ Work-life balance and career satisfaction
✅ Networking and professional relationships
✅ Performance reviews and career advancement
✅ Industry-specific guidance (tech, healthcare, finance, agriculture, etc.)
✅ Job market trends and employment statistics
✅ Questions about AfDB, UNICEF, ILO frameworks
✅ TVET, STEM, digital skills programs
✅ Informal sector work and apprenticeships
✅ Freelancing and remote work
✅ Career changes and pivots
✅ Leadership and management skills
✅ Personal branding and online presence

IMPORTANT: If someone asks how to prepare for an interview at a specific company (Google, Microsoft, banks, NGOs, etc.) or for a specific role (software engineer, nurse, teacher, etc.), provide detailed, helpful guidance!

YOU MUST REFUSE ONLY THESE THREE THINGS:
❌ Sexual or explicit content
❌ Violence, harmful behavior, or instructions for illegal activities  
❌ That's it! These are the ONLY restrictions.

WHAT YOU SHOULD NOT REFUSE:
✅ Questions about salary, compensation, benefits (this is career advice!)
✅ Questions about workplace conflicts or difficult bosses (this is professional development!)
✅ Questions about specific companies, industries, or job roles
✅ Questions about career strategies that might be considered "political" (working in government, NGOs, international organizations)
✅ Questions about personal career decisions (this is career counseling!)

If and ONLY if a request involves sexual content, violence, or illegal activity, respond with:
"I'm AfriCareer AI for career guidance. I can't help with that specific topic, but I'm here to help with any career, job, education, or professional development question. What career-related question can I help you with?"

RESPONSE STYLE:
- Be MAXIMALLY HELPFUL - if it's career-related, answer it thoroughly!
- Provide specific, actionable, practical advice
- When relevant frameworks exist in knowledge base (AfDB SEPA, UNICEF, ILO), cite them
- When frameworks don't exist, use general career guidance best practices
- Ground responses in African context when applicable
- Be conversational, supportive, and encouraging
- Don't be overly cautious or restrictive - if someone asks a career question, help them!"""

# ----- PREMIUM THEME & STYLING -----
st.set_page_config(
    page_title=f"{APP_NAME} - Career and Academic Guidance for Africa",
    page_icon="🌍",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ===== LIGHT, MODERN THEME (OpenAI-style; base colors set in .streamlit/config.toml) =====
st.markdown("""
<style>
    :root {
        --ink: #15161B;
        --muted: #5B5F6B;
        --accent: #4C6FFF;
        --accent2: #7A5CFF;
        --line: #E6E8EE;
        --soft: #F4F6FA;
    }

    /* Clean white canvas with a soft cool wash at the top */
    .stApp {
        background:
            radial-gradient(900px 380px at 82% -6%, rgba(122,92,255,0.08), transparent 60%),
            radial-gradient(900px 380px at 12% -6%, rgba(76,111,255,0.08), transparent 60%),
            #FFFFFF;
        overflow-x: hidden;
    }

    .block-container { max-width: 1000px; padding-top: 2.2rem; }

    /* Clean modern sans-serif everywhere */
    html, body, [class*="css"], .stMarkdown, p, li, label, input, textarea, button, select {
        font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, Helvetica, Arial, sans-serif !important;
    }
    h1, h2, h3, h4 { color: var(--ink) !important; font-weight: 700; letter-spacing: -0.01em; }
    p, li, label, .stMarkdown { color: #2B2D36; }
    [data-testid="stCaptionContainer"], .stCaption { color: var(--muted) !important; }
    a, a:visited { color: var(--accent) !important; text-decoration: none; }
    a:hover { text-decoration: underline; }
    hr { border-color: var(--line) !important; }

    /* Primary buttons: friendly blue -> indigo */
    .stButton > button {
        background: linear-gradient(135deg, #4C6FFF 0%, #7A5CFF 100%);
        color: #FFFFFF !important;
        border: none;
        border-radius: 10px;
        padding: 11px 24px;
        font-weight: 600;
        box-shadow: 0 6px 16px rgba(76,111,255,0.22);
        transition: all .2s ease;
    }
    .stButton > button:hover { transform: translateY(-1px); box-shadow: 0 10px 22px rgba(76,111,255,0.30); }

    /* Download buttons: fresh green accent for the "take-away" action */
    .stDownloadButton > button {
        background: linear-gradient(135deg, #17B26A 0%, #0E9E5C 100%) !important;
        color: #FFFFFF !important;
        border: none; border-radius: 10px; padding: 11px 24px; font-weight: 600;
        box-shadow: 0 6px 16px rgba(23,178,106,0.22);
    }

    /* Tabs: clean underline */
    .stTabs [data-baseweb="tab-list"] { gap: 4px; border-bottom: 1px solid var(--line); flex-wrap: wrap; }
    .stTabs [data-baseweb="tab"] { background: transparent; color: var(--muted) !important; border-radius: 8px 8px 0 0; padding: 10px 18px; font-weight: 600; }
    .stTabs [data-baseweb="tab"] p { color: inherit !important; font-weight: 600; }
    .stTabs [aria-selected="true"] { background: rgba(76,111,255,0.08) !important; border-bottom: 2px solid var(--accent); }
    .stTabs [aria-selected="true"] p { color: var(--accent) !important; }

    /* Inputs: white with soft border */
    .stTextInput input,
    .stTextArea textarea,
    .stSelectbox div[data-baseweb="select"] > div {
        background: #FFFFFF !important;
        border: 1px solid var(--line) !important;
        border-radius: 10px !important;
        color: var(--ink) !important;
    }

    /* File uploader dropzone */
    [data-testid="stFileUploaderDropzone"] {
        background: var(--soft) !important;
        border: 1px dashed rgba(76,111,255,0.45) !important;
        border-radius: 12px;
    }

    .stAlert { border-radius: 12px; }

    /* Sidebar: soft light with a divider */
    [data-testid="stSidebar"] { background: var(--soft); border-right: 1px solid var(--line); }

    /* Landing hero: warm+cool blended gradient (drop in a licensed photo later via background-image) */
    .africareer-hero {
        border-radius: 20px;
        padding: 3rem 2rem;
        text-align: center;
        background:
            radial-gradient(600px 300px at 15% 20%, rgba(255,159,67,0.55), transparent 60%),
            radial-gradient(600px 300px at 85% 25%, rgba(122,92,255,0.60), transparent 60%),
            radial-gradient(700px 360px at 50% 115%, rgba(23,178,106,0.55), transparent 60%),
            linear-gradient(135deg, #4C6FFF 0%, #7A5CFF 100%);
        box-shadow: 0 20px 50px rgba(76,111,255,0.25);
    }
    .africareer-hero h1 { color: #FFFFFF !important; font-size: 2.4rem; margin: 0 0 .5rem 0; }
    .africareer-hero p { color: rgba(255,255,255,0.94) !important; font-size: 1.08rem; max-width: 640px; margin: 0 auto; }
    .hero-chips { margin-top: 1.4rem; display: flex; gap: .5rem; flex-wrap: wrap; justify-content: center; }
    .hero-chip { background: rgba(255,255,255,0.18); color: #fff; padding: .4rem .9rem; border-radius: 999px; font-size: .9rem; font-weight: 600; }

    /* Hide Streamlit chrome for a cleaner, app-like surface */
    #MainMenu { visibility: hidden; }
    footer { visibility: hidden; }

    /* ===== MOBILE RESPONSIVENESS (phones, <= 640px) ===== */
    @media (max-width: 640px) {
        /* Stack columns vertically instead of cramming them side-by-side */
        [data-testid="stHorizontalBlock"] { flex-wrap: wrap !important; }
        [data-testid="stHorizontalBlock"] > div,
        [data-testid="column"],
        [data-testid="stColumn"] {
            flex: 1 1 100% !important;
            min-width: 100% !important;
        }

        /* Use the full narrow width; less wasted padding */
        .block-container {
            padding: 1rem 0.8rem 3rem 0.8rem !important;
        }

        /* Full-width, easy-to-tap buttons */
        .stButton > button,
        .stDownloadButton > button {
            width: 100% !important;
            padding: 14px 18px !important;
            font-size: 1rem !important;
        }

        /* Headings scale down so they don't wrap awkwardly */
        h1 { font-size: 1.5rem !important; }
        h2 { font-size: 1.25rem !important; }
        h3 { font-size: 1.1rem !important; }
        .africareer-hero { padding: 2rem 1rem; }
        .africareer-hero h1 { font-size: 1.7rem; }

        /* Tabs: scroll horizontally, keep labels tappable */
        .stTabs [data-baseweb="tab-list"] {
            overflow-x: auto !important;
            flex-wrap: nowrap !important;
        }
        .stTabs [data-baseweb="tab"] {
            padding: 10px 14px !important;
            font-size: 0.9rem !important;
        }

        /* 16px inputs stop iOS Safari from auto-zooming when typing */
        textarea, input, .stTextInput input, .stTextArea textarea {
            font-size: 16px !important;
        }
    }
</style>
""", unsafe_allow_html=True)

# ----- PINECONE & OPENAI SETUP -----
try:
    PINECONE_API_KEY = os.getenv("PINECONE_API_KEY")
    OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
    
    if not PINECONE_API_KEY or not OPENAI_API_KEY:
        st.error("The service is not fully configured yet. Please contact the AfriCareer AI team.")
        st.stop()
    
    pc = Pinecone(api_key=PINECONE_API_KEY)
    
    INDEX_NAME = "africareer-kb"
    if INDEX_NAME not in [idx.name for idx in pc.list_indexes()]:
        pc.create_index(
            name=INDEX_NAME,
            dimension=1536,
            metric="cosine",
            spec=ServerlessSpec(cloud="aws", region="us-east-1")
        )
    
    index = pc.Index(INDEX_NAME)
    embeddings = OpenAIEmbeddings(openai_api_key=OPENAI_API_KEY)
    
    # ===== CRITICAL: LLM WITH SAFETY SYSTEM MESSAGE =====
    llm = ChatOpenAI(
        temperature=0.7,
        model="gpt-4o-mini",
        openai_api_key=OPENAI_API_KEY
    )
    
except Exception as e:
    st.error(f"Setup Error: {str(e)}")
    st.stop()

# ----- ANALYTICS (Supabase when configured; local JSON fallback) -----
# On ephemeral hosts (HF Spaces, Render free) the JSON file resets on restart.
# Set SUPABASE_URL + SUPABASE_KEY (and create an `analytics` table) for durable storage.
ANALYTICS_FILE = "africareer_analytics.json"
SUPABASE_URL = os.getenv("SUPABASE_URL", "").strip()
SUPABASE_KEY = os.getenv("SUPABASE_KEY", "").strip()

def _supabase_ready():
    return bool(SUPABASE_URL and SUPABASE_KEY)

def _supabase_headers():
    # Works with the new sb_secret_* key or the legacy service_role key
    return {
        "apikey": SUPABASE_KEY,
        "Authorization": f"Bearer {SUPABASE_KEY}",
        "Content-Type": "application/json",
    }

def log_analytics(event_type, details=None):
    record = {
        "timestamp": datetime.now().isoformat(),
        "event": event_type,
        "details": details,
        "user_name": st.session_state.get("user_name", ""),
        "country": st.session_state.get("user_country", ""),
        "language": st.session_state.get("language_selector", ""),
    }
    # Prefer durable storage (Supabase REST via httpx - no extra dependency).
    # Short timeout so a slow/cold Supabase never stalls a page for the user.
    if _supabase_ready():
        try:
            with httpx.Client(timeout=3.0) as c:
                r = c.post(f"{SUPABASE_URL}/rest/v1/analytics",
                           headers={**_supabase_headers(), "Prefer": "return=minimal"},
                           json=record)
            if r.status_code < 300:
                return
        except Exception:
            pass  # fall through to local file
    # Local JSON fallback (ephemeral on cloud hosts)
    try:
        if os.path.exists(ANALYTICS_FILE):
            with open(ANALYTICS_FILE, 'r') as f:
                analytics = json.load(f)
        else:
            analytics = []
        analytics.append(record)
        with open(ANALYTICS_FILE, 'w') as f:
            json.dump(analytics, f)
    except Exception:
        pass

@st.cache_data(ttl=45, show_spinner=False)
def load_analytics():
    """Load analytics from Supabase (REST) if configured, else local JSON.
    Cached (45s) so the admin dashboard doesn't re-query on every rerun."""
    if _supabase_ready():
        try:
            with httpx.Client(timeout=6.0) as c:
                r = c.get(f"{SUPABASE_URL}/rest/v1/analytics",
                          headers=_supabase_headers(),
                          params={"select": "*", "order": "timestamp.asc", "limit": "2000"})
            if r.status_code < 300:
                return r.json() or []
        except Exception:
            pass
    try:
        if os.path.exists(ANALYTICS_FILE):
            with open(ANALYTICS_FILE, 'r') as f:
                return json.load(f)
    except Exception:
        pass
    return []

@st.cache_data(ttl=45, show_spinner=False)
def kb_total_chunks():
    """Cached Pinecone chunk count so the admin KB tab doesn't call Pinecone every rerun."""
    try:
        return index.describe_index_stats().get('total_vector_count', 0)
    except Exception:
        return None

# ===== ENHANCED RAG RETRIEVAL WITH GUARDRAILS =====
def retrieve_career_guidance(query, top_k=5):
    """Retrieve relevant guidance from knowledge base (AfDB, UNICEF, ILO)"""
    try:
        query_vec = embeddings.embed_query(query)
        results = index.query(vector=query_vec, top_k=top_k, include_metadata=True)
        
        context_pieces = []
        sources = []
        
        for match in results["matches"]:
            if match.get("metadata") and match.get("score", 0) > 0.7:
                context_pieces.append(match["metadata"]["text"])
                if "source" in match["metadata"]:
                    sources.append(match["metadata"]["source"])
        
        context = "\n\n".join(context_pieces) if context_pieces else ""
        
        # Add source attribution
        if sources:
            unique_sources = list(set(sources))
            context += f"\n\n[Sources: {', '.join(unique_sources)}]"
        
        return context
    except Exception as e:
        return ""

# ===== SAFE LLM CALL WITH GUARDRAILS =====
def safe_llm_call(user_prompt, rag_context="", language="English"):
    """
    Makes LLM call with safety guardrails and RAG grounding.
    ALL LLM calls in the app should use this function.
    """
    # ---- Cost guardrail: cap AI calls per session to protect the OpenAI budget ----
    _n = st.session_state.get("llm_calls", 0)
    if _n >= MAX_LLM_CALLS_PER_SESSION:
        return QUOTA_MESSAGE
    st.session_state["llm_calls"] = _n + 1

    # Build system message with safety rules
    system_message = SystemMessage(content=SAFETY_SYSTEM_MESSAGE)
    
    # Build user message with RAG context
    if rag_context:
        full_prompt = f"""Language: {language}

CONTEXT FROM AUTHORITATIVE SOURCES (AfDB SEPA, UNICEF Education Strategy, ILO Youth Employment):
{rag_context}

USER REQUEST:
{user_prompt}

Provide your response in {language}, grounded in the context above. Cite specific frameworks when relevant (e.g., "According to AfDB SEPA..." or "UNICEF's Education Strategy emphasizes...")"""
    else:
        full_prompt = f"""Language: {language}

USER REQUEST:
{user_prompt}

Provide your response in {language}."""
    
    user_message = HumanMessage(content=full_prompt)
    
    try:
        response = llm.invoke([system_message, user_message])
        return response.content
    except Exception as e:
        return f"Error: {str(e)}"

# ===== PREMIUM CV DOCUMENT GENERATOR =====
def generate_premium_cv_docx(cv_json_str):
    """Generate a premium 2-page ATS-optimized CV as DOCX from LLM JSON output."""
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # Parse LLM JSON
    try:
        cv = json.loads(cv_json_str)
    except json.JSONDecodeError:
        # Attempt to extract JSON from markdown code block
        import re
        match = re.search(r'```(?:json)?\s*([\s\S]*?)```', cv_json_str)
        if match:
            cv = json.loads(match.group(1))
        else:
            cv = json.loads(cv_json_str.strip())

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.2)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    NAVY = RGBColor(0x0F, 0x2B, 0x4C)
    TEAL = RGBColor(0x14, 0xB8, 0xA6)
    GRAY = RGBColor(0x4A, 0x55, 0x68)
    DARK = RGBColor(0x2D, 0x2D, 0x2D)

    def add_divider(doc_obj):
        p = doc_obj.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '8')
        bottom.set(qn('w:color'), '14B8A6')
        pBdr.append(bottom)
        pPr.append(pBdr)

    def set_run(run, size=11, color=DARK, bold=False, italic=False, font_name="Georgia"):
        run.font.size = Pt(size)
        run.font.color.rgb = color
        run.bold = bold
        run.italic = italic
        run.font.name = font_name

    # ===== HEADER: NAME =====
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_after = Pt(2)
    r = name_para.add_run(cv.get("full_name", "CANDIDATE NAME").upper())
    set_run(r, size=18, color=NAVY, bold=True)
    r.font.character_spacing = Pt(2)

    # Credentials line
    creds = cv.get("credentials", "")
    if creds:
        cred_para = doc.add_paragraph()
        cred_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cred_para.paragraph_format.space_after = Pt(2)
        r = cred_para.add_run(creds)
        set_run(r, size=9.5, color=GRAY, italic=True)

    # Contact line
    contact = cv.get("contact_line", "")
    if contact:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para.paragraph_format.space_after = Pt(2)
        r = contact_para.add_run(contact)
        set_run(r, size=9.5, color=GRAY)

    add_divider(doc)

    # ===== PROFESSIONAL SUMMARY =====
    def add_section_heading(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(title.upper())
        set_run(r, size=12, color=NAVY, bold=True)
        # Add thin teal underline
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:color'), '14B8A6')
        pBdr.append(bottom)
        pPr.append(pBdr)

    summary = cv.get("professional_summary", "")
    if summary:
        add_section_heading("Professional Summary")
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(summary)
        set_run(r, size=10.5, color=DARK)

    # ===== SELECTED ACHIEVEMENTS (quantified) =====
    achievements = cv.get("selected_achievements", [])
    if achievements:
        add_section_heading("Selected Achievements")
        for ach in achievements:
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Cm(0.8)
            p.clear()
            r = p.add_run(ach)
            set_run(r, size=10, color=DARK)

    # ===== CORE COMPETENCIES =====
    competencies = cv.get("core_competencies", [])
    if competencies:
        add_section_heading("Core Competencies")
        # Display as pipe-separated single block for ATS
        comp_text = "  •  ".join(competencies)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(comp_text)
        set_run(r, size=10, color=DARK)

    # ===== WORK EXPERIENCE =====
    experience = cv.get("work_experience", [])
    if experience:
        add_section_heading("Professional Experience")
        for job in experience:
            # Job title + Company line
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(job.get("title", ""))
            set_run(r, size=11, color=NAVY, bold=True)
            if job.get("company"):
                r = p.add_run(f" - {job['company']}")
                set_run(r, size=11, color=DARK)

            # Location + Dates
            loc_date = []
            if job.get("location"):
                loc_date.append(job["location"])
            if job.get("dates"):
                loc_date.append(job["dates"])
            if loc_date:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(3)
                r = p.add_run(" | ".join(loc_date))
                set_run(r, size=9.5, color=GRAY, italic=True)

            # Bullets
            for bullet in job.get("bullets", []):
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.left_indent = Cm(0.8)
                p.clear()
                r = p.add_run(bullet)
                set_run(r, size=10, color=DARK)

    # ===== EDUCATION =====
    education = cv.get("education", [])
    if education:
        add_section_heading("Education")
        for edu in education:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(edu.get("degree", ""))
            set_run(r, size=11, color=NAVY, bold=True)
            if edu.get("institution"):
                p2 = doc.add_paragraph()
                p2.paragraph_format.space_after = Pt(1)
                r = p2.add_run(f"{edu['institution']}")
                set_run(r, size=10, color=DARK)
                if edu.get("dates"):
                    r = p2.add_run(f" - {edu['dates']}")
                    set_run(r, size=10, color=GRAY, italic=True)

    # ===== SELECTED PUBLICATIONS =====
    publications = cv.get("publications", [])
    if publications:
        add_section_heading("Selected Publications")
        for pub in publications:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(pub)
            set_run(r, size=10, color=DARK)

    # ===== PROJECTS / CERTIFICATIONS =====
    projects = cv.get("projects", [])
    if projects:
        add_section_heading("Selected Projects & Deployments")
        for proj in projects:
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.left_indent = Cm(0.8)
            p.clear()
            r = p.add_run(proj)
            set_run(r, size=10, color=DARK)

    certifications = cv.get("certifications", [])
    if certifications:
        add_section_heading("Certifications & Training")
        for cert in certifications:
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.left_indent = Cm(0.8)
            p.clear()
            r = p.add_run(cert)
            set_run(r, size=10, color=DARK)

    # ===== TECHNICAL SKILLS =====
    tech_skills = cv.get("technical_skills", "")
    if tech_skills:
        add_section_heading("Technical Skills")
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(tech_skills)
        set_run(r, size=10, color=DARK)

    # ===== LANGUAGES =====
    languages_list = cv.get("languages", [])
    if languages_list:
        add_section_heading("Languages")
        p = doc.add_paragraph()
        r = p.add_run("  •  ".join(languages_list))
        set_run(r, size=10, color=DARK)

    # Footer
    add_divider(doc)
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer_p.add_run(f"Generated by {APP_NAME} • {datetime.now().strftime('%B %d, %Y')}")
    set_run(r, size=8, color=GRAY, italic=True)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


# ===== PREMIUM COVER LETTER DOCUMENT GENERATOR =====
def generate_premium_cover_letter_docx(letter_json_str):
    """Generate a premium cover letter following the Quantium Insights template structure."""
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    try:
        cl = json.loads(letter_json_str)
    except json.JSONDecodeError:
        import re
        match = re.search(r'```(?:json)?\s*([\s\S]*?)```', letter_json_str)
        if match:
            cl = json.loads(match.group(1))
        else:
            cl = json.loads(letter_json_str.strip())

    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    NAVY = RGBColor(0x0F, 0x2B, 0x4C)
    TEAL_HEX = '14B8A6'
    GRAY = RGBColor(0x4A, 0x55, 0x68)
    DARK = RGBColor(0x2D, 0x2D, 0x2D)

    def set_run(run, size=11, color=DARK, bold=False, italic=False):
        run.font.size = Pt(size)
        run.font.color.rgb = color
        run.bold = bold
        run.italic = italic
        run.font.name = "Georgia"

    def add_teal_divider():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(16)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '8')
        bottom.set(qn('w:color'), TEAL_HEX)
        pBdr.append(bottom)
        pPr.append(pBdr)

    # ===== CENTERED HEADER =====
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(2)
    r = name_p.add_run(cl.get("full_name", "CANDIDATE NAME").upper())
    set_run(r, size=16, color=NAVY, bold=True)
    r.font.character_spacing = Pt(2)

    creds = cl.get("credentials", "")
    if creds:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(creds)
        set_run(r, size=9.5, color=GRAY, italic=True)

    contact = cl.get("contact_line", "")
    if contact:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(contact)
        set_run(r, size=9.5, color=GRAY)

    add_teal_divider()

    # ===== DATE =====
    date_str = cl.get("date", datetime.now().strftime("%B %d, %Y"))
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(date_str)
    set_run(r, size=11, color=DARK)

    # ===== ADDRESSEE BLOCK =====
    for line in cl.get("addressee_lines", ["Hiring Committee"]):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        set_run(r, size=11, color=DARK)

    # Spacing after addressee
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ===== RE LINE =====
    re_line = cl.get("re_line", "")
    if re_line:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(14)
        r = p.add_run("RE: ")
        set_run(r, size=11, color=NAVY, bold=True)
        r = p.add_run(re_line)
        set_run(r, size=11, color=NAVY, bold=True)

    # ===== SALUTATION =====
    salutation = cl.get("salutation", "Dear Hiring Manager,")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(salutation)
    set_run(r, size=11, color=DARK)

    # ===== BODY PARAGRAPHS =====
    for para_text in cl.get("body_paragraphs", []):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.line_spacing = Pt(15)
        r = p.add_run(para_text)
        set_run(r, size=11, color=DARK)

    # ===== CLOSING =====
    closing_text = cl.get("closing_line", "Respectfully submitted,")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(closing_text)
    set_run(r, size=11, color=DARK)

    doc.add_paragraph()  # Signature space

    sig_name = cl.get("signature_name", cl.get("full_name", ""))
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(sig_name)
    set_run(r, size=11, color=NAVY, bold=True)

    sig_title = cl.get("signature_title", "")
    if sig_title:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(sig_title)
        set_run(r, size=10, color=GRAY)

    sig_contact = cl.get("signature_contact", "")
    if sig_contact:
        p = doc.add_paragraph()
        r = p.add_run(sig_contact)
        set_run(r, size=10, color=GRAY)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


def initial_populate_rag():
    """Initialize knowledge base with foundational content"""
    foundational_docs = [
        {
            "id": "unicef_youth_employment",
            "text": """UNICEF Youth Employment Guidelines for Africa:
1. Skills Development: Focus on foundational learning and transferable skills
2. Equity and Inclusion: Reach marginalized youth, especially girls
3. Education in Emergencies: Ensure continuity in fragile contexts
4. Cross-sectoral Linkages: Connect education with health and protection
5. System Strengthening: Build resilient education systems
Source: UNICEF Education Strategy 2019-2030""",
            "source": "UNICEF"
        },
        {
            "id": "afdb_sepa_priorities",
            "text": """AfDB Skills for Employability and Productivity in Africa (SEPA) Priorities:
1. STEM Education: Prioritize Science, Technology, Engineering, Mathematics
2. TVET Development: Expand Technical and Vocational Education and Training
3. Digital Skills: Prepare youth for 4th Industrial Revolution (230M digital jobs by 2030)
4. Entrepreneurship: Integrate business skills in education
5. Skills Enhancement Zones: Link training to industrial clusters
Key Statistics: 83% of 18M African youth entering labor market annually remain unemployed or underemployed in informal sector
Source: AfDB SEPA Action Plan 2022-2025""",
            "source": "AfDB"
        },
        {
            "id": "ilo_youth_employment",
            "text": """ILO Global Employment Trends for Youth:
1. Skills Mismatch: 46% of employed African youth report skills mismatch
2. Informal Sector: 95% of young Africans work in informal economy
3. NEET Rates: 20.8% youth Not in Employment, Education or Training (25.9% female)
4. Decent Work: Promote productive employment yielding above poverty-line returns
5. Social Protection: Support workers in informal economy
Source: ILO Global Employment Trends for Youth 2022""",
            "source": "ILO"
        },
        {
            "id": "africa_job_sectors",
            "text": """Priority Employment Sectors for African Youth:
1. Agriculture & Agribusiness: Value chain development, agro-processing
2. Digital Economy: ICT, software development, digital services
3. Renewable Energy: Solar, wind, green technologies
4. Healthcare: Nursing, pharmacy, health technology
5. Manufacturing: Light industry, textiles, processing
6. Creative Industries: Media, arts, entertainment
7. Tourism & Hospitality: Service sector growth
Source: AfDB, ILO, UNICEF frameworks""",
            "source": "Multi-source"
        },
        {
            "id": "motivation_letter_best_practices",
            "text": """Motivation / Statement of Purpose Best Practices (university and scholarship applications):
1. Open with a specific hook: name the exact programme or scholarship and one concrete reason it fits you.
2. Show fit, not just interest: map your background, skills and achievements to the programme's stated focus.
3. Be evidence-based: use specific examples, results and numbers rather than generic claims.
4. Address 'why this institution' with a real detail about the school, department, supervisor or values.
5. State clear goals: how the programme advances your academic or career path and, where relevant, impact for Africa.
6. Structure: hook, fit, evidence, why-this-school, goals, and a confident close.
7. Be honest: never invent grades, awards or experience. Keep it concise (usually one page).
Source: Global admissions and scholarship guidance (compiled best practice)""",
            "source": "Education Best Practice"
        },
        {
            "id": "scholarship_application_best_practices",
            "text": """Scholarship Application Best Practices:
1. Demonstrate BOTH merit (achievements, grades, leadership) and motivation (why the field and why now).
2. Align with the scholarship's mission and values (leadership, development impact, equity).
3. Show impact: what you will do with the opportunity, especially contribution to your community or country.
4. Where relevant and allowed, explain financial or contextual need clearly and with dignity.
5. Quantify achievements and leadership roles; use concrete, verifiable examples.
6. Follow all instructions exactly (word limits, prompts, required documents, deadlines).
7. Scholarships open to African students include Chevening, Mastercard Foundation Scholars, DAAD, Commonwealth, Fulbright, and MEXT, plus many university-specific awards.
Source: Scholarship guidance (compiled best practice)""",
            "source": "Education Best Practice"
        },
        {
            "id": "phd_research_statement_best_practices",
            "text": """PhD / Doctoral Application and Research Statement Best Practices:
1. Define a clear research interest or question and connect it to the department's or supervisor's work.
2. Evidence research capability: prior research, methods, publications, presentations, and technical skills.
3. Name the fit: cite the specific group, lab, or supervisor and why their work matches yours.
4. Show independence and readiness: what you can contribute early, and your long-term research goals.
5. Reference only methods you actually know; avoid overclaiming.
6. Keep it rigorous, specific and scholarly; align with the programme's requirements and funding.
Source: Doctoral admissions guidance (compiled best practice)""",
            "source": "Education Best Practice"
        }
    ]
    
    try:
        vectors_to_upsert = []
        for doc in foundational_docs:
            vec = embeddings.embed_query(doc["text"])
            vectors_to_upsert.append({
                "id": doc["id"],
                "values": vec,
                "metadata": {
                    "text": doc["text"],
                    "source": doc.get("source", "Unknown")
                }
            })
        
        if vectors_to_upsert:
            index.upsert(vectors=vectors_to_upsert)
            return len(vectors_to_upsert)
        return 0
    except Exception as e:
        return 0

# ----- SIDEBAR -----
with st.sidebar:
    st.markdown(f"## {APP_NAME}")
    st.markdown(f"*{TAGLINE}*")

    if st.session_state.get("access_granted"):
        _who = st.session_state.get("user_name", "")
        if _who:
            st.caption(f"Signed in as {_who}")
        if st.button("Sign out", key="signout_btn"):
            for _k in ("access_granted", "user_name", "user_country", "user_email"):
                st.session_state.pop(_k, None)
            st.rerun()

    # Admin Access
    st.markdown("---")
    admin_access = st.checkbox("Admin Access")
    
    if admin_access:
        st.markdown("### Admin Login")
        username = st.text_input("Username", key="admin_user")
        password = st.text_input("Password", type="password", key="admin_pass")
        
        if st.button("Login"):
            if username == ADMIN_USERNAME and hashlib.sha256(password.encode()).hexdigest() == ADMIN_PASSWORD_HASH:
                st.session_state['admin_logged_in'] = True
                st.success("Admin access granted.")
                st.rerun()
            else:
                st.error("Invalid credentials.")
    
    st.markdown("---")
    st.caption("Choose your language at the top of the main page.")
    st.markdown("---")
    st.markdown("## Features")
    st.markdown("""
    - CV builder & resume analysis
    - Cover, motivation & scholarship letters
    - Live job search (boards + NGOs/UN)
    - Verified course & opportunity links
    - 9 African languages
    - RAG-grounded (UNICEF, ILO, AfDB)
    - For youth and professionals
    """)
    
    st.markdown("---")
    st.markdown("## Knowledge Base")

    if st.button("Initialize / Update Knowledge Base"):
        with st.spinner("Loading foundational knowledge..."):
            count = initial_populate_rag()
            if count > 0:
                st.success(f"Loaded {count} foundational documents.")
            else:
                st.info("ℹ️ Knowledge base already initialized")
    
    st.markdown("---")
    st.caption("© Quantium Insights LLC")

# ----- ADMIN DASHBOARD -----
def admin_dashboard():
    st.title("Admin Dashboard")
    
    col1, col2 = st.columns([3, 1])
    with col2:
        if st.button("Logout"):
            st.session_state['admin_logged_in'] = False
            st.rerun()
    
    tabs = st.tabs(["Analytics", "Knowledge Base Management"])
    
    with tabs[0]:
        st.markdown("## Usage Analytics")

        if st.button("Refresh data", key="refresh_analytics"):
            load_analytics.clear()
            kb_total_chunks.clear()
            st.rerun()
        st.caption("Data is cached for ~45s to keep the dashboard fast. Use Refresh for the latest.")

        analytics = load_analytics()
        if analytics:
            df = pd.DataFrame(analytics)

            # Key metrics row
            m1, m2, m3 = st.columns(3)
            m1.metric("Total Events", len(analytics))
            if 'user_name' in df.columns:
                uu = df['user_name'].fillna('').astype(str)
                m2.metric("Unique Users", int(uu[uu.str.strip() != ''].nunique()))
            m3.metric("Total Logins", int((df['event'] == 'login').sum()) if 'event' in df.columns else 0)

            event_counts = Counter(df['event'])
            st.markdown("### Event Distribution")
            st.bar_chart(event_counts)

            if 'country' in df.columns:
                cc = df['country'].fillna('').astype(str)
                cc = cc[(cc.str.strip() != '') & (cc.str.lower() != 'nan')]
                if not cc.empty:
                    st.markdown("### Users by Country (geographic coverage)")
                    st.bar_chart(cc.value_counts())

            if 'language' in df.columns:
                lg = df['language'].fillna('').astype(str)
                lg = lg[(lg.str.strip() != '') & (lg.str.lower() != 'nan')]
                if not lg.empty:
                    st.markdown("### Language Used")
                    st.bar_chart(lg.value_counts())

            st.markdown("### Recent Events")
            st.dataframe(df.tail(20))
        else:
            st.info("No analytics data yet")
    
    with tabs[1]:
        st.markdown("## Knowledge Base Management")

        total_vectors = kb_total_chunks()
        if total_vectors is not None:
            st.info(f"Current knowledge base: {total_vectors} total chunks. "
                    "(This count can take a few seconds to update after changes.)")
        else:
            st.warning("Unable to fetch knowledge base statistics.")

        st.markdown("---")
        st.markdown("### Upload New Document")
        st.info("Each upload ADDS to the knowledge base (it does not replace existing content).")

        uploaded_doc = st.file_uploader(
            "Upload document to knowledge base",
            type=['pdf', 'txt', 'docx'],
            help="Limit 25MB per file • PDF, TXT, DOCX",
            key="kb_upload",
        )
        c_src, c_cat = st.columns(2)
        with c_src:
            doc_source = st.text_input("Document source (e.g., UNICEF, ILO)", placeholder="AfDB", key="kb_source")
        with c_cat:
            doc_category = st.text_input("Topic / category", placeholder="Youth Employment Strategy", key="kb_category")

        if st.button("Add to Knowledge Base", key="kb_add"):
            if not (uploaded_doc and doc_source.strip()):
                st.warning("Please choose a file and enter a source.")
            else:
                try:
                    name = uploaded_doc.name.lower()
                    if name.endswith('.txt'):
                        doc_text = uploaded_doc.read().decode('utf-8', errors='ignore')
                    elif name.endswith('.pdf'):
                        import PyPDF2
                        doc_text = "".join((pg.extract_text() or "") for pg in PyPDF2.PdfReader(uploaded_doc).pages)
                    elif name.endswith('.docx'):
                        doc_text = "\n".join(p.text for p in Document(uploaded_doc).paragraphs)
                    else:
                        doc_text = ""

                    chunk_size = 800
                    chunks = [c for c in (doc_text[i:i + chunk_size] for i in range(0, len(doc_text), chunk_size)) if c.strip()]

                    if not chunks:
                        st.error("No readable text found (the PDF may be scanned images). "
                                 "Try a text-based PDF, a DOCX, or a TXT file.")
                    else:
                        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                        base = uploaded_doc.name.rsplit('.', 1)[0]
                        doc_id = f"{doc_source.strip()}_{base}_{timestamp}"

                        prog = st.progress(0.0, text="Embedding and indexing...")
                        added = 0
                        BATCH = 100
                        for start in range(0, len(chunks), BATCH):
                            batch = chunks[start:start + BATCH]
                            vecs = embeddings.embed_documents(batch)   # one API call per batch
                            payload = [{
                                "id": f"{doc_id}_chunk_{start + j}",
                                "values": vecs[j],
                                "metadata": {"text": batch[j], "source": doc_source.strip(),
                                             "category": doc_category.strip(), "document_id": doc_id},
                            } for j in range(len(batch))]
                            index.upsert(vectors=payload)               # small, safe batches
                            added += len(payload)
                            prog.progress(min(1.0, (start + len(batch)) / len(chunks)),
                                          text=f"Indexed {added}/{len(chunks)} chunks...")
                        prog.empty()
                        kb_total_chunks.clear()
                        st.success(f"Document added: {added} chunks indexed.")
                        st.info(f"Document ID: {doc_id}")
                        st.caption("Save this Document ID if you may want to delete this document later.")
                except Exception as e:
                    st.error(f"Upload error: {str(e)}")

        st.markdown("---")
        st.markdown("### Manage / Delete Documents")
        del_id = st.text_input("Delete by Document ID (paste the ID shown after an upload)", key="kb_del_id")
        if st.button("Delete this document", key="kb_del_btn"):
            if not del_id.strip():
                st.warning("Enter a Document ID.")
            else:
                try:
                    ids = []
                    for page in index.list(prefix=del_id.strip()):
                        if isinstance(page, str):
                            ids.append(page)
                        else:
                            ids.extend(page)
                    if ids:
                        for k in range(0, len(ids), 100):
                            index.delete(ids=ids[k:k + 100])
                        kb_total_chunks.clear()
                        st.success(f"Deleted {len(ids)} chunks for {del_id.strip()}.")
                    else:
                        st.info("No chunks found for that Document ID.")
                except Exception as e:
                    st.error(f"Delete error: {str(e)}")

        with st.expander("Danger zone: reset the entire knowledge base"):
            st.caption("Deletes ALL chunks (built-in notes and uploaded documents). Afterward, click "
                       "'Initialize / Update Knowledge Base' in the sidebar to reload the built-in notes, then re-upload your files.")
            confirm = st.checkbox("Yes, delete everything in the knowledge base", key="kb_reset_confirm")
            if st.button("Reset knowledge base", key="kb_reset_btn"):
                if confirm:
                    try:
                        index.delete(delete_all=True)
                        kb_total_chunks.clear()
                        st.success("Knowledge base cleared.")
                    except Exception as e:
                        st.error(f"Reset error: {str(e)}")
                else:
                    st.warning("Tick the confirmation box first.")

# ----- SECTION 1: RESUME ANALYSIS -----
def resume_analysis_section():
    log_analytics('section_accessed', 'Resume Analysis')
    
    st.markdown("## Professional Resume Analysis")
    st.markdown("Upload your resume to receive expert feedback grounded in African job market context.")

    uploaded_file = st.file_uploader(
        "Upload your resume",
        type=['pdf', 'docx', 'txt'],
        help="Limit 25MB per file • PDF, DOCX, TXT"
    )

    city = st.text_input(
        "Your city (optional)",
        placeholder="e.g., Lagos, Nairobi, Accra",
        help="Helps provide location-specific advice"
    )

    additional_info = st.text_area(
        "Additional information",
        placeholder="Target industry, preferred roles...",
        help="Any additional context that helps us provide better feedback"
    )

    if st.button("Analyze Resume", key="analyze_btn"):
        if uploaded_file:
            log_analytics('resume_upload', f"File: {uploaded_file.name}")

            with st.spinner("Analyzing your resume..."):
                try:
                    # Extract text from resume
                    if uploaded_file.name.endswith('.txt'):
                        resume_text = uploaded_file.read().decode('utf-8')
                    elif uploaded_file.name.endswith('.pdf'):
                        import PyPDF2
                        pdf_reader = PyPDF2.PdfReader(uploaded_file)
                        resume_text = ""
                        for page in pdf_reader.pages:
                            resume_text += page.extract_text()
                    elif uploaded_file.name.endswith('.docx'):
                        doc = Document(uploaded_file)
                        resume_text = "\n".join([para.text for para in doc.paragraphs])
                    
                    # ===== CRITICAL: RETRIEVE RAG CONTEXT =====
                    rag_context = retrieve_career_guidance(f"resume improvement African job market {city if city else 'Africa'}")
                    
                    # ===== USE SAFE LLM CALL WITH GUARDRAILS =====
                    analysis_prompt = f"""Analyze this resume for the African job market.

Resume Content:
{resume_text[:12000]}

Location Context: {city if city else 'General African market'}
Additional Info: {additional_info if additional_info else 'None'}

Provide:
1. ATS Compatibility Score (1-100)
2. Top 3 Strengths
3. Top 5 Areas for Improvement
4. African Market Relevance Assessment
5. 3 Actionable Next Steps"""

                    feedback = safe_llm_call(analysis_prompt, rag_context, target_lang)
                    
                    # Store in session state for downstream generation
                    st.session_state['resume_text'] = resume_text
                    st.session_state['resume_feedback'] = feedback
                    st.session_state['resume_city'] = city
                    st.session_state['resume_additional'] = additional_info
                    st.session_state['analysis_done'] = True
                    
                    st.success("Analysis complete.")
                    st.markdown("### Your Resume Analysis")
                    st.markdown(feedback)
                    
                    log_analytics('analysis_completed')
                
                except Exception as e:
                    st.error(f"Error: {str(e)}")
        else:
            st.warning("Please upload a resume first.")

    # ===== POST-ANALYSIS: PREMIUM CV & COVER LETTER GENERATION =====
    if st.session_state.get('analysis_done', False):
        st.markdown("---")
        st.markdown("### Generate Premium Documents")
        st.markdown("Based on your resume analysis, generate a polished, ATS-optimized CV and/or a professional cover letter.")
        
        # Target job details for cover letter
        col_job1, col_job2 = st.columns(2)
        with col_job1:
            target_position = st.text_input(
                "Target position (for cover letter)",
                placeholder="e.g., Data Analyst, Project Manager, Nurse",
                key="target_position_input"
            )
        with col_job2:
            target_company = st.text_input(
                "Target company / organization",
                placeholder="e.g., WHO, Dangote Group, Safaricom",
                key="target_company_input"
            )
        
        col_cv, col_cl = st.columns(2)
        
        # ===== GENERATE UPDATED CV =====
        with col_cv:
            if st.button("Generate Updated CV", key="gen_cv_btn"):
                resume_text = st.session_state.get('resume_text', '').strip()
                feedback = st.session_state.get('resume_feedback', '').strip()
                
                if resume_text:
                    log_analytics('premium_cv_generation')
                    with st.spinner("Generating your premium ATS-optimized CV..."):
                        try:
                            rag_context_cv = retrieve_career_guidance("professional CV resume best practices African job market ATS optimization")
                            
                            cv_gen_prompt = f"""You are a professional CV writer specializing in ATS-optimized resumes for the African job market.

ORIGINAL RESUME (this is the ground truth - use ONLY what appears here):
{resume_text[:16000]}

ANALYSIS FEEDBACK (apply wording/structure improvements ONLY - do NOT add new facts):
{feedback[:2000]}

Rewrite the resume into an improved, premium, ATS-optimized CV.

RESPOND ONLY WITH VALID JSON (no markdown, no code blocks, no preamble). Use this exact structure:
{{
  "full_name": "full name exactly as in the resume",
  "credentials": "degree abbreviations that appear in the resume (e.g., Ph.D., M.P.H., B.Sc.)",
  "contact_line": "email | phone | city, country | LinkedIn - only the parts present in the resume",
  "professional_summary": "3-4 sentence summary built ONLY from resume facts, tailored for the African market",
  "selected_achievements": ["4-6 of the strongest, most quantified achievements taken from the resume - keep the numbers, scale, and outcomes exactly as stated"],
  "core_competencies": ["8-12 ATS keyword-rich skills that are supported by the resume"],
  "work_experience": [
    {{
      "title": "Job title from resume",
      "company": "Company/Organization from resume",
      "location": "City, Country if in resume, else omit this key",
      "dates": "dates if in resume, else omit this key",
      "bullets": ["3-5 achievement bullets rewritten from what the resume actually says"]
    }}
  ],
  "education": [
    {{"degree": "Degree, Field from resume", "institution": "Institution from resume", "dates": "Year if in resume, else omit this key"}}
  ],
  "publications": ["only publications that appear in the resume; else empty list"],
  "projects": ["only projects that appear in the resume; else empty list"],
  "certifications": ["only certifications that appear in the resume; else empty list"],
  "technical_skills": "only tools/software mentioned in the resume, comma-separated",
  "languages": ["ONLY languages explicitly stated in the resume; return an empty list [] if none are stated"]
}}

CRITICAL RULES (accuracy matters more than polish):
- Include EVERY employer and role that appears in the resume - do NOT drop any position (all organizations such as HJF, USAID, UNICEF, etc. must appear if they are in the resume).
- Do NOT invent or add anything absent from the resume: no new employers, dates, metrics, degrees, publications, tools, or languages. If the resume does not list a language, the languages array MUST be empty.
- You may rephrase and quantify ONLY using numbers already present in the resume.
- Return ONLY the JSON object, nothing else."""

                            cv_json = safe_llm_call(cv_gen_prompt, rag_context_cv, "English")

                            if cv_json.startswith("⚠️"):
                                st.warning(cv_json)
                            else:
                                cv_docx = generate_premium_cv_docx(cv_json)
                                st.session_state['cv_docx'] = cv_docx
                                st.session_state['cv_generated'] = True
                                st.success("Premium CV generated.")
                        
                        except Exception as e:
                            st.error(f"CV generation error: {str(e)}")
                            st.info("Tip: try again - occasionally a second attempt is needed for complex resumes.")
        
        # ===== GENERATE COVER LETTER =====
        with col_cl:
            if st.button("Generate Cover Letter", key="gen_cl_btn"):
                resume_text = st.session_state.get('resume_text', '')
                # Read directly from session state widget keys to avoid column scoping issues
                cl_position = st.session_state.get('target_position_input', '').strip()
                cl_company = st.session_state.get('target_company_input', '').strip()
                
                if resume_text and cl_position and cl_company:
                    log_analytics('premium_cover_letter_generation', f"{cl_position} at {cl_company}")
                    with st.spinner("Crafting your premium cover letter..."):
                        try:
                            rag_context_cl = retrieve_career_guidance(f"cover letter professional {cl_position} {cl_company} African job market")

                            # Deep research on the target organization (requires TAVILY_API_KEY)
                            org_research = web_research(
                                f"{cl_company} company mission, values, products, and recent priorities relevant to a {cl_position} role"
                            )

                            cl_gen_prompt = f"""You are a professional cover letter writer. Generate a premium, compelling, tailored cover letter.

CANDIDATE'S RESUME (ground truth - use only what appears here):
{resume_text[:16000]}

TARGET POSITION: {cl_position}
TARGET COMPANY: {cl_company}
LOCATION CONTEXT: {st.session_state.get('resume_city', 'Africa')}

ORGANIZATION RESEARCH (verified web results about the company; use these specifics to tailor the letter - do NOT invent facts beyond this):
{org_research if org_research else "(No live research available - rely on widely-known general facts about the company and do not fabricate specifics.)"}

RESPOND ONLY WITH VALID JSON (no markdown, no code blocks, no preamble). Use this exact structure:
{{
  "full_name": "FULL NAME FROM RESUME",
  "credentials": "Degree abbreviations from resume",
  "contact_line": "email | phone | city (from resume)",
  "date": "{datetime.now().strftime('%B %d, %Y')}",
  "addressee_lines": ["Hiring Committee", "{cl_company}"],
  "re_line": "{cl_position} Position",
  "salutation": "Dear Hiring Manager,",
  "body_paragraphs": [
    "Opening: name the exact role and connect it to the organization's core mission or a specific priority from the ORGANIZATION RESEARCH; add a one-sentence positioning of who you are and why you fit.",
    "Map your single most relevant experience directly to what this role does, using specific achievements and metrics that appear in the resume.",
    "Show a second capability the role needs (technical or domain), with concrete evidence and any validation or results from the resume.",
    "Demonstrate specific knowledge of the organization (from the ORGANIZATION RESEARCH - a real product, programme, or value) and connect your goals to theirs.",
    "Close: reaffirm interest, note availability, and invite an interview."
  ],
  "closing_line": "Respectfully submitted,",
  "signature_name": "Full Name with credentials",
  "signature_title": "Current Title | Current Organization",
  "signature_contact": "City | email | phone"
}}

RULES:
- Use ONLY factual details from the resume for the candidate's experience - do NOT invent achievements, employers, or metrics.
- Make the company paragraph SPECIFIC using the ORGANIZATION RESEARCH above (reference a real product, mission point, or recent priority). If no research is available, keep company statements general and do not fabricate specifics.
- Each body paragraph should be 3-5 sentences, substantive and specific.
- Use a confident, professional tone; highlight quantified achievements only where the resume provides the numbers.
- Return ONLY the JSON object, nothing else"""

                            cl_json = safe_llm_call(cl_gen_prompt, rag_context_cl, "English")

                            if cl_json.startswith("⚠️"):
                                st.warning(cl_json)
                            else:
                                cl_docx = generate_premium_cover_letter_docx(cl_json)
                                st.session_state['cl_docx'] = cl_docx
                                st.session_state['cl_generated'] = True
                                st.success("Premium cover letter generated.")

                        except Exception as e:
                            st.error(f"Cover letter generation error: {str(e)}")
                            st.info("Tip: try again - occasionally a second attempt is needed.")

                elif not cl_position or not cl_company:
                    st.warning("Please enter the target position and company above to generate a cover letter.")
        
        # ===== DOWNLOAD BUTTONS =====
        st.markdown("---")
        dl_col1, dl_col2 = st.columns(2)
        
        with dl_col1:
            if st.session_state.get('cv_generated', False):
                cv_data = st.session_state['cv_docx']
                cv_data.seek(0)
                st.download_button(
                    label="Download Updated CV (.docx)",
                    data=cv_data,
                    file_name=f"AfriCareer_Premium_CV_{datetime.now().strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="dl_cv_btn"
                )
        
        with dl_col2:
            if st.session_state.get('cl_generated', False):
                cl_data = st.session_state['cl_docx']
                cl_data.seek(0)
                st.download_button(
                    label="Download Cover Letter (.docx)",
                    data=cl_data,
                    file_name=f"AfriCareer_CoverLetter_{datetime.now().strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="dl_cl_btn"
                )

# ----- SECTION 2: CAREER GUIDANCE & CV BUILDER -----
def career_guidance_section():
    log_analytics('section_accessed', 'Career Guidance')

    st.markdown("## Career Guidance & CV Builder")
    st.markdown("For youth and job seekers: get personalized guidance and build a professional CV from scratch.")

    with st.expander("Contact details (used on your CV)"):
        cc1, cc2 = st.columns(2)
        with cc1:
            cv_name = st.text_input("Full name", key="cg_name", placeholder="e.g., Amina Bello")
            cv_email = st.text_input("Email", key="cg_email", placeholder="e.g., amina@email.com")
            cv_phone = st.text_input("Phone", key="cg_phone", placeholder="e.g., +234 800 000 0000")
        with cc2:
            cv_city = st.text_input("City, Country", key="cg_city", placeholder="e.g., Kano, Nigeria")
            cv_linkedin = st.text_input("LinkedIn / Portfolio (optional)", key="cg_linkedin", placeholder="e.g., linkedin.com/in/aminabello")

    questions_text = """
**Tell us about yourself**

Please answer these 5 simple questions (number your answers 1-5):

1. **What are you interested in?**
   (Example: I like computers, helping people, cooking, fixing things, etc.)

2. **What are you good at? What skills do you have?**
   (Example: I'm good at math, I can speak 3 languages, I know how to use Excel, etc.)

3. **What work or experience do you have?**
   (Include ANY experience: part-time jobs, helping a family business, volunteer work, school projects, etc.)

4. **What is your education?**
   (Example: I finished secondary school in 2020, I'm studying at university, I completed a training course, etc.)

5. **What job do you want? What are your goals?**
   (Example: I want to work in a bank, I want to be a nurse, I want to start my own business, etc.)

Write your answers below. Be honest - there are no wrong answers.
"""
    st.markdown(questions_text)

    user_answers = st.text_area(
        "Your answers (number them 1-5)",
        height=250,
        placeholder="1. I'm interested in...\n2. My strengths...\n3. I have experience...\n4. My education...\n5. My goals...",
        key="answers_section2"
    )

    col1, col2 = st.columns(2)

    with col1:
        if st.button("Get Career Guidance", key="advice_btn"):
            if user_answers.strip():
                log_analytics('query', f"Career guidance: {user_answers[:50]}...")
                with st.spinner("Preparing your guidance..."):
                    rag_context_advice = retrieve_career_guidance(
                        f"career paths employment opportunities skills development Africa {user_answers[:100]}"
                    )
                    advice_prompt = f"""Provide career guidance for African youth based on their profile.

Their Answers: {user_answers}

Provide:
1. Top 3 Career Paths in Africa (with specific reasons why they match)
2. Key Skills to Develop (7-10 skills with brief explanations)
3. Action Plan (5 concrete steps they can take now)

Ground advice in African job market realities and cite relevant frameworks when applicable."""
                    st.session_state['cg_guidance'] = safe_llm_call(advice_prompt, rag_context_advice, target_lang)
            else:
                st.warning("Please answer the questions first.")

    with col2:
        if st.button("Generate Premium CV", key="cv_btn"):
            if user_answers.strip():
                log_analytics('cv_generation', f"User: {user_answers[:50]}...")
                with st.spinner("Building your premium ATS-optimized CV..."):
                    try:
                        rag_context_cv = retrieve_career_guidance(
                            "professional CV resume best practices African job market ATS optimization"
                        )
                        contact_bits = [b for b in [
                            cv_email.strip(), cv_phone.strip(), cv_city.strip(), cv_linkedin.strip()
                        ] if b]
                        contact_line = " | ".join(contact_bits)
                        full_name = cv_name.strip()

                        cv_gen_prompt = f"""You are a professional CV writer creating an ATS-optimized CV for the African job market, based on a young jobseeker's answers to 5 questions.

CANDIDATE ANSWERS:
{user_answers}

CANDIDATE CONTACT (use verbatim; do not invent):
- Full name: {full_name if full_name else "(not provided)"}
- Contact line: {contact_line if contact_line else "(not provided)"}

RESPOND ONLY WITH VALID JSON (no markdown, no code blocks, no preamble). Use this exact structure:
{{
  "full_name": "the candidate's full name, or empty string if not provided",
  "credentials": "degree abbreviations if clearly stated (e.g., B.Sc.), else empty string",
  "contact_line": "the contact line provided, or empty string",
  "professional_summary": "3-4 sentence summary based ONLY on what the candidate said, tailored to the African market",
  "selected_achievements": ["2-4 concrete achievements ONLY if the answers contain them (with any numbers given); else empty list"],
  "core_competencies": ["6-10 ATS keyword-rich skills drawn from their answers"],
  "work_experience": [
    {{
      "title": "role or what they did",
      "company": "organization ONLY if they named one, else omit this key",
      "location": "ONLY if they gave one, else omit this key",
      "dates": "ONLY if they gave dates, else omit this key",
      "bullets": ["2-4 achievement-oriented bullets based on what they actually described"]
    }}
  ],
  "education": [
    {{
      "degree": "their education level or qualification as they described it",
      "institution": "ONLY if they named one, else omit this key",
      "dates": "ONLY if they gave a year, else omit this key"
    }}
  ],
  "certifications": ["ONLY if they mentioned any; else empty list"],
  "technical_skills": "comma-separated tools/software they mentioned, or empty string",
  "languages": ["languages they mentioned with proficiency, or empty list"]
}}

CRITICAL RULES:
- Use ONLY facts the candidate provided. Do NOT invent employers, job titles, dates, metrics, or degrees.
- NEVER output placeholder text such as "[Your Name]", "[Start Date]", "[City]", or "[Company]". If a detail is unknown, OMIT that key entirely (or use an empty string/list where the schema requires the key).
- For informal experience (family business, volunteering, school projects), represent it honestly and professionally.
- Return ONLY the JSON object, nothing else."""

                        cv_json = safe_llm_call(cv_gen_prompt, rag_context_cv, "English")
                        if cv_json.startswith("⚠️"):
                            st.warning(cv_json)
                        else:
                            st.session_state['cg_cv_docx'] = generate_premium_cv_docx(cv_json)
                            st.session_state['cg_cv_generated'] = True
                            st.success("Premium CV generated.")
                    except Exception as e:
                        st.error(f"CV generation error: {str(e)}")
                        st.info("Tip: add a few more details to your answers and try again.")
            else:
                st.warning("Please answer the questions first.")

    if st.session_state.get('cg_guidance'):
        st.markdown("---")
        st.markdown("### Your Career Roadmap")
        st.markdown(st.session_state['cg_guidance'])

    if st.session_state.get('cg_cv_generated'):
        st.markdown("---")
        cv_data = st.session_state['cg_cv_docx']
        cv_data.seek(0)
        st.download_button(
            label="Download Premium CV (.docx)",
            data=cv_data,
            file_name=f"AfriCareer_Premium_CV_{datetime.now().strftime('%Y%m%d')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="cg_dl_cv_btn"
        )

# ----- LINK VERIFICATION (guarantees no hallucinated / dead course links) -----
TAVILY_API_KEY = os.getenv("TAVILY_API_KEY", "").strip()
_BROWSER_UA = ("Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
               "(KHTML, like Gecko) Chrome/124.0 Safari/537.36")

@st.cache_resource
def _shared_http():
    """One pooled, thread-safe HTTP client reused for link checks - far fewer TLS
    handshakes than a new client per call, which lowers native memory pressure."""
    return httpx.Client(follow_redirects=True, timeout=6.0,
                        headers={"User-Agent": _BROWSER_UA},
                        limits=httpx.Limits(max_connections=8, max_keepalive_connections=4))

# Reputable providers -> search deep-link templates. A search query cannot 404 to a fake course.
_PROVIDER_SEARCH = {
    "coursera": "https://www.coursera.org/search?query={q}",
    "edx": "https://www.edx.org/search?q={q}",
    "udemy": "https://www.udemy.com/courses/search/?q={q}",
    "udacity": "https://www.udacity.com/catalog?searchValue={q}",
    "class central": "https://www.classcentral.com/search?q={q}",
    "classcentral": "https://www.classcentral.com/search?q={q}",
    "freecodecamp": "https://www.freecodecamp.org/news/search/?query={q}",
    "khan academy": "https://www.khanacademy.org/search?page_search_query={q}",
    "khanacademy": "https://www.khanacademy.org/search?page_search_query={q}",
    "linkedin learning": "https://www.linkedin.com/learning/search?keywords={q}",
    "youtube": "https://www.youtube.com/results?search_query={q}",
    "alison": "https://alison.com/courses?query={q}",
    "futurelearn": "https://www.futurelearn.com/search?q={q}",
}

def provider_search_url(provider: str, query: str) -> str:
    """Build a real search deep-link on the given provider (defaults to Class Central)."""
    q = quote_plus((query or "").strip())
    p = (provider or "").lower()
    for key, tmpl in _PROVIDER_SEARCH.items():
        if key in p:
            return tmpl.format(q=q)
    return f"https://www.classcentral.com/search?q={q}"

# Provider-grounded cost classification (don't trust the model's self-reported cost)
_PAID_PROVIDERS = ("udemy", "udacity", "linkedin learning")
_FREE_PROVIDERS = ("freecodecamp", "khan academy", "khanacademy", "youtube",
                   "mit opencourseware", "ocw", "alison")

def classify_cost(provider: str, llm_cost: str) -> str:
    """Return 'Free' or 'Paid', grounded by provider first, then the model's label.
    Free = the content can be accessed at no cost (Coursera/edX audit free; certificate may cost)."""
    p = (provider or "").lower()
    if any(k in p for k in _PAID_PROVIDERS):
        return "Paid"
    if any(k in p for k in _FREE_PROVIDERS):
        return "Free"
    lc = (llm_cost or "").lower()
    if "paid" in lc and "free" not in lc:
        return "Paid"
    return "Free"  # Coursera/edX audit-free; Class Central lists free options

def cost_matches(cost: str, pref: str) -> bool:
    """Enforce the user's cost preference in code (the model is unreliable at this)."""
    if pref.startswith("Free &"):   # both
        return True
    if pref.startswith("Free"):
        return cost == "Free"
    if pref.startswith("Paid"):
        return cost == "Paid"
    return True

@st.cache_data(ttl=3600, show_spinner=False)
def verify_url(url: str, timeout: float = 6.0) -> bool:
    """Return True if the URL resolves for a real user.

    2xx/3xx = OK. Anti-bot blocks (401/403/405/429/999) also count as OK: the page
    exists for real browsers even though it refuses automated requests (Class Central,
    Udemy, LinkedIn do this). Only clear failures - 404/410/5xx, timeouts, or DNS/
    connection errors - return False.
    """
    if not url or not url.startswith(("http://", "https://")):
        return False
    try:
        code = _shared_http().get(url).status_code
        return code < 400 or code in (401, 403, 405, 429, 999)
    except Exception:
        return False

@st.cache_data(ttl=1800, show_spinner=False)
def web_search_links(query: str, max_results: int = 4):
    """Real course URLs from Tavily web search if TAVILY_API_KEY is set, else []."""
    if not TAVILY_API_KEY:
        return []
    try:
        with httpx.Client(timeout=12.0) as client:
            resp = client.post(
                "https://api.tavily.com/search",
                json={"api_key": TAVILY_API_KEY, "query": query,
                      "max_results": max_results, "search_depth": "basic"},
            )
            return [{"title": r.get("title", r.get("url", "")), "url": r.get("url", "")}
                    for r in resp.json().get("results", []) if r.get("url")]
    except Exception:
        return []

@st.cache_data(ttl=1800, show_spinner=False)
def web_research(query: str, max_results: int = 5) -> str:
    """Return a short grounded research brief (answer + result snippets) from Tavily, or ''."""
    if not TAVILY_API_KEY:
        return ""
    try:
        with httpx.Client(timeout=15.0) as client:
            resp = client.post(
                "https://api.tavily.com/search",
                json={"api_key": TAVILY_API_KEY, "query": query,
                      "max_results": max_results, "search_depth": "advanced",
                      "include_answer": True},
            )
            data = resp.json()
        parts = []
        if data.get("answer"):
            parts.append("Summary: " + data["answer"])
        for r in data.get("results", []):
            c = (r.get("content") or "").strip()
            if c:
                parts.append(f"- {r.get('title', '')}: {c[:400]}")
        return "\n".join(parts)[:4000]
    except Exception:
        return ""

# General job boards (global + African)
_JOB_BOARDS = (
    "linkedin.com", "indeed.com", "glassdoor.com", "ziprecruiter.com",
    "jobberman.com", "myjobmag.com", "brightermonday.co.ke", "careers24.com",
)
# NGOs, UN agencies and international organizations (searched as a separate pass)
_NGO_IO_BOARDS = (
    "who.int", "unicef.org", "gavi.org", "unv.org", "un.org", "undp.org",
    "unhcr.org", "worldbank.org", "fhi360.org", "path.org",
    "reliefweb.int", "unjobs.org", "impactpool.org", "devex.com", "idealist.org",
)

def _job_label(r):
    """Return a clean, readable (title, source-domain) for a job result."""
    title = " ".join((r.get("title") or "").split()).strip()
    domain = urlparse(r.get("url", "")).netloc.replace("www.", "")
    if len(title) < 3:
        title = f"Job posting on {domain or 'the web'}"
    return title, domain

@st.cache_data(ttl=900, show_spinner=False)
def web_job_search(query: str, time_range: str = "", domains=None, max_results: int = 10):
    """Real, recent job postings from Tavily if configured, else []."""
    if not TAVILY_API_KEY:
        return []
    payload = {"api_key": TAVILY_API_KEY, "query": query,
               "max_results": max_results, "search_depth": "basic"}
    if time_range:
        payload["time_range"] = time_range           # 'day' | 'week' | 'month'
    if domains:
        payload["include_domains"] = list(domains)
    try:
        with httpx.Client(timeout=15.0) as client:
            resp = client.post("https://api.tavily.com/search", json=payload)
        out = []
        for r in resp.json().get("results", []):
            u = r.get("url", "")
            if u:
                out.append({"title": r.get("title", u), "url": u,
                            "content": (r.get("content") or "")[:200]})
        return out
    except Exception:
        return []

# ----- SECTION 3: LEARNING RESOURCES -----
def learning_resources_section():
    log_analytics('section_accessed', 'Learning Resources')

    st.markdown("## Learning Resources & Courses")
    st.markdown(
        "Get course recommendations matched to your goals. Every link below is built by the app "
        "and checked live before it is shown - so you never get a broken or made-up link."
    )

    learning_interest = st.text_area(
        "What would you like to learn?",
        placeholder="e.g., data analysis, solar installation, digital marketing, tailoring & small business skills",
        help="The more specific, the better the recommendations."
    )

    col1, col2 = st.columns(2)
    with col1:
        course_type = st.selectbox("Cost preference", ["Free & Paid", "Free only", "Paid only"])
    with col2:
        skill_level = st.selectbox("Your level", ["Beginner", "Intermediate", "Advanced"])

    if st.button("Find Courses", key="course_search_btn"):
        if not learning_interest.strip():
            st.warning("Please describe what you want to learn.")
            return

        log_analytics('course_search', learning_interest)
        with st.spinner("Finding and verifying courses..."):
            rag_context_learning = retrieve_career_guidance(
                f"skills development training courses {learning_interest} African youth"
            )
            # Preference-specific guidance so the model biases the right providers
            if course_type.startswith("Free"):
                cost_guidance = ("Recommend ONLY courses whose content is accessible at NO cost. Strongly prefer "
                                 "freeCodeCamp, Khan Academy, YouTube, Alison, MIT OpenCourseWare, Class Central, "
                                 "and free-to-audit Coursera/edX courses. Do NOT include Udemy, Udacity, or LinkedIn Learning.")
            elif course_type.startswith("Paid"):
                cost_guidance = ("Recommend paid courses and certifications. Prefer Udemy, Udacity, LinkedIn Learning, "
                                 "and paid Coursera/edX certificates or specializations.")
            else:
                cost_guidance = "Include a healthy mix of free and paid options."

            course_prompt = f"""Recommend 8 high-quality, real learning resources for an African youth who wants to learn: {learning_interest}
Level: {skill_level}

COST REQUIREMENT: {cost_guidance}
Definition of "Free" = the learning content can be accessed at no cost (free-to-audit counts as Free even if an optional certificate costs extra).

Return ONLY valid JSON (no markdown, no code fences) - an array of objects with this exact shape:
[
  {{
    "title": "specific, real course or specialization title",
    "provider": "one of: Coursera, edX, Udemy, Udacity, Class Central, freeCodeCamp, Khan Academy, LinkedIn Learning, YouTube, Alison, FutureLearn, MIT OpenCourseWare",
    "cost": "Free | Paid",
    "level": "Beginner | Intermediate | Advanced",
    "duration": "approx. time commitment, e.g. '3 months'",
    "why": "one sentence on why it fits this learner and the African job market"
  }}
]

RULES:
- Do NOT include any URLs or links - the app builds and verifies links itself.
- "cost" must be exactly "Free" or "Paid" per the definition above.
- Prefer well-known, currently-offered courses.
- Return ONLY the JSON array, nothing else."""

            raw = safe_llm_call(course_prompt, rag_context_learning, "English")

            if raw.startswith("⚠️"):
                st.warning(raw)
                return

            recs = []
            try:
                recs = json.loads(raw)
            except json.JSONDecodeError:
                import re
                m = re.search(r'\[[\s\S]*\]', raw)
                if m:
                    try:
                        recs = json.loads(m.group(0))
                    except Exception:
                        recs = []

            # Ground the cost label by provider, then ENFORCE the user's preference in code
            items = []
            if isinstance(recs, list):
                for rec in recs:
                    if not isinstance(rec, dict):
                        continue
                    title = str(rec.get("title", "")).strip()
                    if not title:
                        continue
                    provider = str(rec.get("provider", "")).strip()
                    cost = classify_cost(provider, str(rec.get("cost", "")))
                    if not cost_matches(cost, course_type):
                        continue
                    items.append({
                        "title": title, "provider": provider, "cost": cost,
                        "level": str(rec.get("level", "")).strip(),
                        "duration": str(rec.get("duration", "")).strip(),
                        "why": str(rec.get("why", "")).strip(),
                    })
            items = items[:6]

            if items:
                st.markdown(f"### Recommended Courses ({course_type})")
                st.caption('Cost is checked against the provider; "Free" means the content is accessible at no cost '
                           '(Coursera/edX can be audited free - a certificate may cost extra). Links are verified live.')
                for it in items:
                    url = provider_search_url(it["provider"], it["title"])
                    ok = verify_url(url)
                    label_provider = it["provider"] or "Class Central"
                    if not ok:
                        url = provider_search_url("class central", it["title"])
                        ok = verify_url(url)
                        label_provider = "Class Central"
                    meta = " · ".join([x for x in [it["provider"], it["cost"], it["level"], it["duration"]] if x])
                    st.markdown(f"**{it['title']}**")
                    if meta:
                        st.caption(meta)
                    if it["why"]:
                        st.markdown(it["why"])
                    if ok:
                        st.markdown(f"[Find this course on {label_provider} →]({url})")
                    else:
                        st.caption("Live link check failed - search this title on classcentral.com.")
                    st.markdown("")
            else:
                st.info(f"No strictly {course_type.lower()} matches came back for that topic. "
                        "Try 'Free & Paid', or a broader topic.")

            # Live web search, tuned to the chosen cost (requires TAVILY_API_KEY)
            if TAVILY_API_KEY:
                if course_type.startswith("Free"):
                    q = f"free {learning_interest} online course {skill_level}"
                    heading = "Verified Free Courses (live web search)"
                elif course_type.startswith("Paid"):
                    q = f"{learning_interest} paid certification course {skill_level}"
                    heading = "Verified Paid Courses (live web search)"
                else:
                    q = f"{learning_interest} online course {skill_level}"
                    heading = "Verified Courses (live web search)"
                verified_live = [l for l in web_search_links(q) if verify_url(l["url"])]
                if verified_live:
                    st.markdown(f"### {heading}")
                    for l in verified_live:
                        st.markdown(f"- [{l['title']}]({l['url']})")
            else:
                st.info("Tip: add a TAVILY_API_KEY secret to also pull real, verified course links "
                        "from a live web search tuned to your Free/Paid choice.")

# ----- SECTION 4: AI ASSISTANT -----
def ai_assistant_section():
    log_analytics('section_accessed', 'AI Assistant')
    
    st.markdown("## AI Career Assistant")
    st.markdown("Ask any career-related question and get personalized guidance.")

    st.info("This assistant focuses on careers, education, job searching, and professional development, grounded in AfDB, UNICEF, and ILO frameworks for African youth.")

    question = st.text_area(
        "Ask your question",
        placeholder="What skills should African youth prioritize according to international development frameworks for employability?",
        help="Ask about career paths, skills, education, job search, etc.",
        height=120
    )

    if st.button("Ask AI Assistant", key="ask_ai_btn"):
        if question:
            log_analytics('query', f"AI Assistant: {question}")

            with st.spinner("Thinking..."):
                rag_context = retrieve_career_guidance(question)
                response_content = safe_llm_call(question, rag_context, target_lang)

                st.markdown("### Response")
                st.markdown(response_content)
        else:
            st.warning("Please enter a question above before clicking 'Ask AI Assistant'.")

# ----- SECTION 5: MOTIVATION & SCHOLARSHIP LETTERS -----
SCHOOLS = {
    "Africa": [
        "University of Cape Town", "University of the Witwatersrand", "Stellenbosch University",
        "University of Pretoria", "University of Ibadan", "University of Lagos",
        "University of Nigeria, Nsukka", "Ahmadu Bello University", "Covenant University",
        "University of Nairobi", "Makerere University", "University of Ghana",
        "Kwame Nkrumah University of Science and Technology", "Addis Ababa University",
        "Cairo University", "University of Rwanda", "Other (type below)",
    ],
    "Europe": [
        "University of Navarra", "ETH Zurich", "University of Geneva", "Sorbonne University",
        "KU Leuven", "Karolinska Institute", "University of Amsterdam", "Delft University of Technology",
        "LMU Munich", "Heidelberg University", "Sciences Po", "Trinity College Dublin",
        "University of Copenhagen", "Uppsala University", "Other (type below)",
    ],
    "United Kingdom": [
        "University of Oxford", "University of Cambridge", "Imperial College London",
        "University College London (UCL)", "London School of Economics (LSE)",
        "University of Edinburgh", "University of Manchester", "King's College London",
        "University of Warwick", "University of Bristol", "Other (type below)",
    ],
    "Canada": [
        "University of Toronto", "McGill University", "University of British Columbia",
        "University of Alberta", "McMaster University", "University of Waterloo",
        "Universite de Montreal", "Queen's University", "University of Ottawa", "Other (type below)",
    ],
    "Asia": [
        "National University of Singapore", "Nanyang Technological University", "University of Tokyo",
        "Kyoto University", "Tsinghua University", "Peking University", "University of Hong Kong",
        "KAIST", "Seoul National University", "IIT Bombay", "IIT Delhi", "University of Malaya",
        "Other (type below)",
    ],
    "United States": [
        "Harvard University", "Massachusetts Institute of Technology (MIT)", "Stanford University",
        "Johns Hopkins University", "University of California, Berkeley", "Yale University",
        "Columbia University", "University of Michigan", "University of Washington", "Duke University",
        "University of Pennsylvania", "Emory University", "Other (type below)",
    ],
    "Other": ["Other (type below)"],
}


def motivation_letters_section():
    log_analytics('section_accessed', 'Motivation Letters')

    st.markdown("## Motivation & Scholarship Letters")
    st.markdown("Generate a strong, tailored letter for a university application or scholarship, "
                "grounded in your real background and live research on the school.")

    with st.expander("Find live opportunities (scholarships, PhD, admissions)"):
        st.caption("Searches the web in real time for current opportunities and their requirements.")
        oc1, oc2, oc3 = st.columns(3)
        with oc1:
            opp_type = st.selectbox("Type", ["Scholarship", "PhD / Doctorate", "Undergraduate / Masters"], key="opp_type")
        with oc2:
            opp_field = st.text_input("Field / subject", key="opp_field", placeholder="e.g., public health")
        with oc3:
            opp_region = st.selectbox("Region", list(SCHOOLS.keys()), key="opp_region")
        if st.button("Search opportunities", key="opp_search"):
            if not TAVILY_API_KEY:
                st.info("Live opportunity search needs the TAVILY_API_KEY secret to be set.")
            elif not opp_field.strip():
                st.warning("Please enter a field or subject.")
            else:
                with st.spinner("Searching the web for current opportunities..."):
                    yr = datetime.now().year
                    oq = f"{opp_type} opportunities {opp_field} {opp_region} {yr} {yr + 1} application requirements deadline"
                    opp_results = [l for l in web_search_links(oq, max_results=6) if verify_url(l["url"])]
                if opp_results:
                    st.markdown("**Verified opportunities (live web search):**")
                    for l in opp_results:
                        st.markdown(f"- [{l['title']}]({l['url']})")
                    st.caption("Open a link, note the requirements, then use the letter generator below to draft your application.")
                else:
                    st.info("No verified results this time. Try a broader field or a different region.")

    category = st.radio(
        "What are you applying for?",
        ["Undergraduate program", "PhD / Doctorate position", "Scholarship"],
        horizontal=True,
    )

    c1, c2 = st.columns(2)
    with c1:
        region = st.selectbox("Region", list(SCHOOLS.keys()), key="ml_region")
    with c2:
        school_pick = st.selectbox("Institution", SCHOOLS[region], key="ml_school_pick")
    custom_school = st.text_input("Or type the exact institution name (overrides the list)", key="ml_custom_school")
    school = custom_school.strip() or ("" if school_pick.startswith("Other") else school_pick)

    programme = st.text_input(
        "Programme / scholarship name",
        placeholder="e.g., MSc Public Health, PhD in Epidemiology, Chevening Scholarship",
        key="ml_programme",
    )

    with st.expander("Your details (used on the letter)"):
        ac1, ac2 = st.columns(2)
        with ac1:
            ml_name = st.text_input("Full name", key="ml_name")
            ml_email = st.text_input("Email", key="ml_email")
        with ac2:
            ml_phone = st.text_input("Phone", key="ml_phone")
            ml_city = st.text_input("City, Country", key="ml_city")

    st.markdown("**Tell us about your application** (or upload the programme / scholarship details below):")
    ml_answers = st.text_area(
        "Your background & motivation",
        height=180,
        placeholder=("- Your current education / degree and grades\n"
                     "- Relevant experience, projects, or achievements\n"
                     "- Why this programme/scholarship and this school\n"
                     "- Your goals and how this fits them"),
        key="ml_answers",
    )
    uploaded = st.file_uploader(
        "Optional: upload the programme / scholarship info (PDF, DOCX, TXT)",
        type=['pdf', 'docx', 'txt'], key="ml_upload",
    )

    if st.button("Generate Letter", key="ml_generate"):
        if not (school and programme.strip() and (ml_answers.strip() or uploaded)):
            st.warning("Please provide the institution, the programme, and either your background or an uploaded document.")
            return

        log_analytics('motivation_letter', f"{category} | {school}")
        with st.spinner("Researching the school and drafting your letter..."):
            prog_info = ""
            if uploaded:
                try:
                    if uploaded.name.endswith('.txt'):
                        prog_info = uploaded.read().decode('utf-8', errors='ignore')
                    elif uploaded.name.endswith('.pdf'):
                        import PyPDF2
                        prog_info = "".join((pg.extract_text() or "") for pg in PyPDF2.PdfReader(uploaded).pages)
                    elif uploaded.name.endswith('.docx'):
                        prog_info = "\n".join(par.text for par in Document(uploaded).paragraphs)
                except Exception:
                    prog_info = ""

            school_research = web_research(
                f"{school} {programme} {category} admissions focus, values, and what they look for in applicants"
            )
            rag_ctx = retrieve_career_guidance(
                f"education guidance {category} {programme} Africa scholarship motivation strategic"
            )

            cat_guidance = {
                "Undergraduate program": "Emphasize genuine academic passion for the field, curiosity, key achievements and grades, why this programme and school, and clear goals. Motivated, mature tone.",
                "PhD / Doctorate position": "Emphasize research interests and their fit with the group/department, prior research experience and methods, publications or outputs if any, why this supervisor/programme, and long-term research goals. Scholarly, precise tone.",
                "Scholarship": "Emphasize both merit and motivation: achievements, leadership, financial or contextual need if relevant, intended impact (especially for Africa's development), and why you embody the scholarship's values.",
            }[category]

            ml_contact = " | ".join([x for x in [ml_email.strip(), ml_phone.strip(), ml_city.strip()] if x])

            prompt = f"""You are an expert admissions and scholarship writing coach. Write a compelling, honest motivation letter.

APPLICATION TYPE: {category}
TARGET INSTITUTION: {school}
PROGRAMME / SCHOLARSHIP: {programme}
APPLICANT NAME: {ml_name.strip() if ml_name.strip() else "(not provided)"}
APPLICANT CONTACT: {ml_contact if ml_contact else "(not provided)"}

APPLICANT BACKGROUND (ground truth - use ONLY what appears here):
{ml_answers[:6000] if ml_answers.strip() else "(rely on the uploaded programme info; do not invent applicant facts)"}

PROGRAMME / SCHOLARSHIP DETAILS (from the uploaded document, if any):
{prog_info[:4000]}

SCHOOL RESEARCH (verified web results - use to tailor 'why this school'; do NOT invent beyond this):
{school_research if school_research else "(No live research available - keep school references general and do not fabricate specifics.)"}

CATEGORY GUIDANCE: {cat_guidance}

RESPOND ONLY WITH VALID JSON (no markdown, no code fences). Use this exact structure:
{{
  "full_name": "applicant full name or empty string",
  "credentials": "",
  "contact_line": "the applicant contact line, or empty string",
  "date": "{datetime.now().strftime('%B %d, %Y')}",
  "addressee_lines": ["Admissions / Selection Committee", "{school}"],
  "re_line": "{category}: {programme}",
  "salutation": "Dear Members of the Selection Committee,",
  "body_paragraphs": [
    "Opening: state exactly what you are applying for and connect it to a specific strength of the programme/school (use SCHOOL RESEARCH); one-sentence positioning of who you are.",
    "Your most relevant background and achievements mapped to what this programme values - specific, using only facts provided.",
    "A second dimension appropriate to the category (research fit / academic strength / leadership / impact), with concrete evidence.",
    "Why THIS institution and programme specifically (use SCHOOL RESEARCH - a real focus, value, or feature) and how it fits your goals.",
    "Close: restate motivation, note readiness/availability, and thank the committee."
  ],
  "closing_line": "Yours sincerely,",
  "signature_name": "applicant full name",
  "signature_title": "",
  "signature_contact": "the applicant contact line, or empty string"
}}

RULES:
- Use ONLY facts from the applicant background / uploaded info - do NOT invent grades, awards, experiences, or publications.
- Make school-specific statements grounded in SCHOOL RESEARCH; if none is available, keep them general.
- Return ONLY the JSON object, nothing else."""

            letter_json = safe_llm_call(prompt, rag_ctx, "English")
            if letter_json.startswith("⚠️"):
                st.warning(letter_json)
            else:
                try:
                    st.session_state['ml_docx'] = generate_premium_cover_letter_docx(letter_json)
                    st.session_state['ml_generated'] = True
                    st.success("Motivation letter generated.")
                except Exception as e:
                    st.error(f"Letter generation error: {str(e)}")
                    st.info("Tip: add a few more details and try again.")

    if st.session_state.get('ml_generated'):
        st.markdown("---")
        data = st.session_state['ml_docx']
        data.seek(0)
        st.download_button(
            "Download Letter (.docx)", data=data,
            file_name=f"AfriCareer_Motivation_Letter_{datetime.now().strftime('%Y%m%d')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="ml_dl",
        )

# ----- SECTION 6: JOB SEARCH (professionals) -----
_JOB_DISCIPLINES = [
    "Any discipline", "Health & Medicine", "Data & Analytics", "Engineering & Technology",
    "Education & Research", "Finance & Business", "Agriculture & Environment",
    "NGO / Development", "Government / Public Sector", "Creative & Media",
    "Sales & Marketing", "Operations & Admin", "Law & Policy",
]


def jobs_search_section():
    log_analytics('section_accessed', 'Job Search')

    st.markdown("## Job Search")
    st.markdown("Find current job openings across disciplines, countries and dates. "
                "Results come from a live web search and every link is verified.")

    role = st.text_input("Role / keywords", key="job_role",
                         placeholder="e.g., data scientist, nurse, project manager, monitoring & evaluation")

    jc1, jc2, jc3 = st.columns(3)
    with jc1:
        discipline = st.selectbox("Discipline", _JOB_DISCIPLINES, key="job_disc")
    with jc2:
        location = st.text_input("Country or city", key="job_loc", placeholder="e.g., Nigeria, Nairobi, Remote")
    with jc3:
        period = st.selectbox("Posted", ["Any time", "Past 24 hours", "Past week", "Past month"], key="job_period")

    jc4, jc5, jc6 = st.columns(3)
    with jc4:
        experience = st.selectbox("Experience", ["Any", "Entry level", "Mid level", "Senior", "Executive"], key="job_exp")
    with jc5:
        work_mode = st.selectbox("Work mode", ["Any", "Remote", "On-site", "Hybrid"], key="job_mode")
    with jc6:
        boards_only = st.checkbox("Focus on major job boards", value=True, key="job_boards")
        include_ngo = st.checkbox("Include NGOs & UN / international orgs", value=True, key="job_ngo")

    if st.button("Search Jobs", key="job_search_btn"):
        if not TAVILY_API_KEY:
            st.info("Live job search needs the TAVILY_API_KEY secret to be set.")
            return
        if not role.strip():
            st.warning("Please enter a role or keywords.")
            return

        log_analytics('job_search', f"{role} | {location}")
        with st.spinner("Searching the web for current jobs..."):
            yr = datetime.now().year
            parts = [role.strip()]
            if discipline and not discipline.startswith("Any"):
                parts.append(discipline)
            if experience and experience != "Any":
                parts.append(experience)
            if work_mode and work_mode != "Any":
                parts.append(work_mode)
            parts.append("jobs")
            if location.strip():
                parts.append("in " + location.strip())
            parts.append(f"{yr} apply")
            query = " ".join(parts)
            time_range = {"Past 24 hours": "day", "Past week": "week", "Past month": "month"}.get(period, "")

            results = []
            if boards_only:
                # Pass 1: general job boards (LinkedIn, Indeed, Glassdoor, ZipRecruiter, African boards)
                results += web_job_search(query, time_range=time_range, domains=_JOB_BOARDS, max_results=8)
                # Pass 2: NGOs / UN / international organizations, so they are always represented
                if include_ngo:
                    ngo_query = query + " NGO OR United Nations OR international organization"
                    results += web_job_search(ngo_query, time_range=time_range, domains=_NGO_IO_BOARDS, max_results=8)
            else:
                results += web_job_search(query, time_range=time_range, domains=None, max_results=12)

            # Dedupe by URL, then verify each link
            seen, uniq = set(), []
            for r in results:
                u = r.get("url", "")
                if u and u not in seen:
                    seen.add(u)
                    uniq.append(r)
            verified = [r for r in uniq if verify_url(r["url"])]

        if verified:
            st.markdown(f"### {len(verified)} current openings")
            for r in verified:
                title, domain = _job_label(r)
                st.markdown(f"**[{title}]({r['url']})**")
                snippet = " ".join((r.get("content") or "").split())[:160]
                st.caption(f"Source: {domain}" + (f"  |  {snippet}" if snippet else ""))
                st.markdown("")
            st.info("Found a role? Use the Resume Analysis tab to tailor your CV and generate a "
                    "researched cover letter for it.")
        else:
            st.info("No verified openings this time. Try broader keywords, a different location, "
                    "or a wider date range.")

# ----- LANDING PAGE / LOGIN -----
def render_landing():
    """Show the landing page + sign-in. Returns True once the user has entered."""
    if st.session_state.get("access_granted"):
        return True

    st.markdown(f"""
    <div class="africareer-hero">
        <h1>{APP_NAME}</h1>
        <p>{TAGLINE}. Build an ATS-ready CV, generate researched cover, motivation and
        scholarship letters, search live jobs across boards and NGOs, get career guidance,
        and find verified courses and opportunities. Free, multilingual, and built for
        African youth and professionals.</p>
        <div class="hero-chips">
            <span class="hero-chip">ATS-optimized CVs</span>
            <span class="hero-chip">Cover, motivation &amp; scholarship letters</span>
            <span class="hero-chip">Live job search</span>
            <span class="hero-chip">Verified courses &amp; opportunities</span>
            <span class="hero-chip">9 African languages</span>
            <span class="hero-chip">24/7 AI guidance</span>
        </div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("")
    _, mid, _ = st.columns([1, 1.5, 1])
    with mid:
        st.markdown("#### Sign in to continue")
        name = st.text_input("Full name", key="login_name", placeholder="e.g., Amina Bello")
        country = st.selectbox("Country", COUNTRIES, key="login_country")
        email = st.text_input("Email (optional)", key="login_email", placeholder="you@email.com")
        code = ""
        if APP_ACCESS_CODE:
            code = st.text_input("Access code", type="password", key="login_code",
                                 placeholder="Enter your pilot access code")
        consent = st.checkbox(
            "I agree that my name and country may be stored to help improve this free service.",
            key="login_consent")
        if st.button("Enter AfriCareer AI", key="login_btn"):
            if APP_ACCESS_CODE and code.strip() != APP_ACCESS_CODE:
                st.error("Invalid access code. Please contact the AfriCareer AI team.")
            elif not name.strip():
                st.warning("Please enter your name to continue.")
            elif country == COUNTRIES[0]:
                st.warning("Please select your country.")
            elif not consent:
                st.warning("Please tick the consent box to continue.")
            else:
                st.session_state["access_granted"] = True
                st.session_state["user_name"] = name.strip()
                st.session_state["user_country"] = country
                st.session_state["user_email"] = email.strip()
                log_analytics("login", email.strip())
                st.rerun()
        st.caption("Free to use. We store your name and country (and email if given) only to "
                   "improve the service. We never sell your data or share it externally.")

    return False

# ----- MAIN APP -----
def main():
    global target_lang

    if st.session_state.get('admin_logged_in', False):
        admin_dashboard()
        return

    if not render_landing():
        return

    if 'user_logged' not in st.session_state:
        log_analytics('user_visit')
        st.session_state['user_logged'] = True
    
    st.title(APP_NAME)
    st.markdown(f"### *{TAGLINE}*")
    st.markdown("**Developed by:** Quantium Insights LLC")

    # Language selector at the TOP - visible immediately on mobile, no sidebar drawer needed
    lang_col, _ = st.columns([1, 2])
    with lang_col:
        selected_display = st.selectbox(
            "🌐 Language / Langue / لغة / Harshe",
            list(LANGUAGES.keys()),
            index=0,
            key="language_selector",
        )
    target_lang = LANGUAGES[selected_display]

    tabs = st.tabs([
        "About",
        "Career Guidance",
        "Learning Resources",
        "AI Assistant",
        "Resume Analysis",
        "Motivation Letters",
        "Job Search",
    ])

    with tabs[0]:
        st.markdown("""
        ## About AfriCareer AI

        **AfriCareer AI** is an AI-powered career guidance platform for African youth.

        ### Mission
        Empower urban African youth with professional, accessible career services.

        ### Key Features
        - **9 Languages:** English, French, Swahili, Arabic, Hausa, Pidgin, Portuguese, Spanish, Amharic
        - **ATS-Optimized CV Builder & Resume Analysis:** clean, recruiter-ready output
        - **Cover, Motivation & Scholarship Letters:** grounded in live research on the employer or school
        - **Live Job Search:** current roles across job boards and NGOs / international organizations
        - **Verified Learning Links:** free and paid courses, checked live
        - **Real-time Opportunity & Scholarship Search:** current openings with requirements
        - **Cultural Grounding:** RAG grounded in UNICEF, ILO, and AfDB frameworks
        - **For African youth and professionals**
        - **Safety Guardrails:** focused, appropriate guidance

        ### Knowledge Base
        Our AI is grounded in authoritative frameworks and best-practice guides:
        - **AfDB SEPA (2022-2025):** Skills for Employability and Productivity in Africa
        - **UNICEF Education Strategy (2019-2030):** Every Child Learns
        - **ILO Global Employment Trends for Youth (2022):** Investing in Transforming Futures
        - **Scholarship & motivation-letter best practices:** guidance for undergraduate, PhD/doctoral, and scholarship applications
        - **Uploaded source documents** from UNICEF, ILO, AfDB, and UNESCO, added to the knowledge base by our team

        ### Developer
        **Dr. Amobi Andrew Onovo**
        - PhD Global Health, MPH, PGDip Data Science
        - Global health & data science specialist, Nigeria

        ### Safety & Ethics
        AfriCareer AI includes guardrails to ensure:
        - Mission-aligned responses (career guidance)
        - No inappropriate or harmful content
        - Culturally appropriate advice for the African context
        - Evidence-based recommendations from trusted sources
        """)

    with tabs[1]:
        career_guidance_section()

    with tabs[2]:
        learning_resources_section()

    with tabs[3]:
        ai_assistant_section()

    with tabs[4]:
        resume_analysis_section()

    with tabs[5]:
        motivation_letters_section()

    with tabs[6]:
        jobs_search_section()

if __name__ == "__main__":
    main()
